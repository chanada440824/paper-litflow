"""export 命令: 将摘抄 Markdown 目录导出为 Word 文档 (宋体、层级标题、置信度配色)。

结构兼容两种布局:
- 嵌套: md_root/2.1 xxx/2.1.1 xxx/文献.md
- 扁平: md_root/2.1/文献.md  (extract-v2 产出)
"""

import json
import os
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor


def set_font_to_songti(paragraph):
    """将段落所有 run 的字体设置为宋体。"""
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_heading_font_to_songti(doc, level):
    """设置指定级别的标题样式为宋体。"""
    style = doc.styles[f"Heading {level}"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def confidence_rgb(conf: str):
    """置信度 -> 字体颜色。高=绿, 中=橙, 低=灰, 未知=黑。"""
    c = (conf or "").strip()
    if c == "高":
        return RGBColor(0x00, 0xB0, 0x50)
    if c == "中":
        return RGBColor(0xED, 0x7D, 0x31)
    if c == "低":
        return RGBColor(0x80, 0x80, 0x80)
    return RGBColor(0x00, 0x00, 0x00)


_CONF_RE = re.compile(r"置信度[**：:].*?\[?(高|中|低)\]?")


def extract_confidence(text: str) -> str:
    """从 用途/出处 行提取置信度标签。"""
    m = _CONF_RE.search(text or "")
    return m.group(1) if m else ""


def parse_md_file(md_path: str):
    """解析摘抄 Markdown, 返回 [(原文, 用途, 置信度), ...]。

    支持两种格式:
    1. 引用块格式: '> 原文' + '**用途**：...', 条目间以 '---' 分隔
    2. 列表格式: '- **原文内容**：...' + '- **用途**：...'
    """
    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    blocks = re.split(r"\n---\s*\n", content) if "---" in content else [content]
    entries = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if re.search(r"^>\s", block, re.MULTILINE):
            match = re.search(r"^>\s*(.*?)\n\s*(\*\*.*?)$", block, re.MULTILINE)
            if match:
                quote = match.group(1).strip()
                meta = match.group(2).strip()
                usage = meta.split("**用途**：", 1)[1].strip() if "**用途**：" in meta else meta
                entries.append((quote, usage, extract_confidence(meta)))
        elif "- **原文内容**：" in block:
            entries.extend(parse_list_style(block))
    return entries


def parse_list_style(text: str):
    """逐行扫描列表式摘抄: '- **原文内容**：' 与 '- **用途**：' 配对, 允许空行分隔。"""
    entries = []
    original = None
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("- **原文内容**："):
            original = line.replace("- **原文内容**：", "").strip()
        elif line.startswith("- **用途**："):
            if original:
                usage = line.replace("- **用途**：", "").strip()
                entries.append((original, usage, extract_confidence(usage)))
            original = None
    return entries


def extract_author_year_title(filename_base: str):
    """从文件名 (去扩展名) 提取 (作者, 年份, 标题)。格式: '作者 - 2020 - 标题_28-77'"""
    base = re.sub(r"_\d+-\d+$", "", filename_base)
    base = re.sub(r"\(\d+\)$", "", base).strip()
    parts = base.split(" - ")
    if len(parts) >= 3:
        return parts[0].strip(), parts[1].strip(), " - ".join(parts[2:]).strip()
    elif len(parts) == 2:
        return parts[0].strip(), "", parts[1].strip()
    return "", "", base


_NUM_CN = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def _emit_file(doc, md_file, idx):
    base = os.path.basename(md_file).replace(".md", "")
    base = re.sub(r"\(\d+\)", "", base).strip()
    author, year, pure_title = extract_author_year_title(base)
    label = f"文献{_NUM_CN[idx - 1]}" if idx <= len(_NUM_CN) else f"文献{idx}"
    if author and year:
        heading_text = f"{label}·{author} ({year})·{pure_title}"
    elif author:
        heading_text = f"{label}·{author}·{pure_title}"
    else:
        heading_text = f"{label}·{pure_title}"
    doc.add_heading(heading_text, level=3)

    entries = parse_md_file(os.path.join(md_file))
    if not entries:
        doc.add_paragraph("（本节未提取到有效摘抄）")
    for original, usage, conf in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(original)
        run.italic = True
        run.font.color.rgb = confidence_rgb(conf)
        set_font_to_songti(p)
        p_usage = doc.add_paragraph()
        p_usage.add_run("用途：").bold = True
        add_md_bold_runs(p_usage, usage)
        set_font_to_songti(p_usage)
        doc.add_paragraph("_" * 50)
    doc.add_paragraph()


def add_md_bold_runs(paragraph, text: str):
    """把 '**加粗**' 段渲染为加粗 run, 去掉裸星号。"""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        run.bold = (i % 2 == 1)


def generate_word_from_md(md_root: str, output_docx: str, sections=None):
    doc = Document()
    sections = sections or {}

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for level in range(1, 4):
        set_heading_font_to_songti(doc, level)

    # 收集顶层 2.x 目录, 兼容 2.x / 2.x.y / 嵌套 2.x>2.x.y 三种布局
    top_dirs = sorted(
        d for d in os.listdir(md_root)
        if os.path.isdir(os.path.join(md_root, d)) and re.match(r"^2\.\d", d)
    )
    chapters = {}
    for d in top_dirs:
        code = d.split(" ", 1)[0]
        path = os.path.join(md_root, d)
        if code.count(".") == 1:
            chapters.setdefault(code, {"path": path, "subs": [], "title": d})
        else:
            parent = code.rsplit(".", 1)[0]
            ch = chapters.setdefault(parent, {"path": None, "subs": [], "title": parent})
            ch["subs"].append((code, path, d))
    for code, ch in chapters.items():
        ch["subs"].sort(key=lambda x: x[0])
        if ch["path"]:
            for d in sorted(os.listdir(ch["path"])):
                p = os.path.join(ch["path"], d)
                if os.path.isdir(p) and re.match(r"^2\.\d+\.\d+", d):
                    scode = d.split(" ", 1)[0]
                    if not any(s[0] == scode for s in ch["subs"]):
                        ch["subs"].append((scode, p, d))
        ch["subs"].sort(key=lambda x: x[0])

    for code in sorted(chapters):
        ch = chapters[code]
        doc.add_heading(sections.get(code, ch["title"]), level=1)
        direct_mds = sorted(f for f in os.listdir(ch["path"]) if f.endswith(".md")) if ch["path"] else []
        if ch["subs"]:
            for scode, sdir, sname in ch["subs"]:
                doc.add_heading(sections.get(scode, sname), level=2)
                mds = sorted(f for f in os.listdir(sdir) if f.endswith(".md"))
                for i, mf in enumerate(mds, 1):
                    _emit_file(doc, os.path.join(sdir, mf), i)
        elif direct_mds:
            for i, mf in enumerate(direct_mds, 1):
                _emit_file(doc, os.path.join(ch["path"], mf), i)
        else:
            print(f"警告: {code} 无内容")

    doc.save(output_docx)
    print(f"Word 文档已生成: {output_docx}")


def run(args) -> int:
    if not os.path.exists(args.md_root):
        print(f"根目录不存在: {args.md_root}")
        return 1
    sections = {}
    if getattr(args, "sections_file", None):
        with open(args.sections_file, encoding="utf-8-sig") as f:
            sections = json.load(f)
    generate_word_from_md(args.md_root, args.output, sections)
    return 0
